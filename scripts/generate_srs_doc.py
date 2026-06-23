from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BODY_FONT = "宋体"
HEADING_FONT = "黑体"
EN_FONT = "Times New Roman"
BODY_SIZE = Pt(12)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click to update field."

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(placeholder)
    run._r.append(fld_char_end)


def apply_run_font(run, east_asia: str = BODY_FONT, ascii_font: str = EN_FONT, size: Pt = BODY_SIZE) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = size


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = EN_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = BODY_SIZE

    for style_name, size in (("Heading 1", Pt(16)), ("Heading 2", Pt(14)), ("Heading 3", Pt(12))):
        style = document.styles[style_name]
        style.font.name = EN_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT)
        style.font.size = size
        style.font.bold = True

    title = document.styles["Title"]
    title.font.name = EN_FONT
    title._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT)
    title.font.size = Pt(20)
    title.font.bold = True

    subtitle = document.styles["Subtitle"]
    subtitle.font.name = EN_FONT
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    subtitle.font.size = Pt(12)
    subtitle.font.bold = False


def set_page_margins(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.style = "Title"
    p.add_run("Software Requirements Specification")

    for line in (
        "Project: 校园超速监控系统",
        "Published on: 2026-06-21",
        "Version: 1.0",
    ):
        para = document.add_paragraph()
        para.style = "Subtitle"
        para.add_run(line)

    document.add_page_break()

    toc_heading = document.add_paragraph()
    toc_heading.style = "Heading 1"
    toc_heading.add_run("Table of Contents")
    add_toc(document.add_paragraph())
    document.add_page_break()


def add_code_paragraph(document: Document, text: str) -> None:
    para = document.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9.5)


def add_image(document: Document, image_path: Path, caption: str) -> None:
    width = Inches(6.2)
    if image_path.name in {"system-context.png", "use-case.png"}:
        width = Inches(5.9)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=width)

    caption_para = document.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_para.add_run(caption)
    apply_run_font(caption_run, east_asia=BODY_FONT, ascii_font=EN_FONT, size=Pt(10.5))


def add_normal_paragraph(document: Document, text: str) -> None:
    stripped = text.lstrip()
    numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
    if numbered:
        para = document.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.28)
        para.paragraph_format.first_line_indent = Inches(-0.2)
        run = para.add_run(f"{numbered.group(1)}. {numbered.group(2)}")
        apply_run_font(run)
        return

    dashed = re.match(r"^-\s+(.*)$", stripped)
    if dashed:
        para = document.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.28)
        para.paragraph_format.first_line_indent = Inches(-0.18)
        run = para.add_run(f"• {dashed.group(1)}")
        apply_run_font(run)
        return

    para = document.add_paragraph()
    run = para.add_run(text)
    apply_run_font(run)


def markdown_to_docx(markdown_path: Path, output_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_styles(document)
    set_page_margins(document)
    add_cover(document)

    in_code = False
    for line in lines:
        if line.startswith("```"):
            in_code = not in_code
            continue

        if in_code:
            add_code_paragraph(document, line)
            continue

        if not line.strip():
            document.add_paragraph("")
            continue

        if line.startswith("# "):
            if line[2:].strip() == "Software Requirements Specification":
                continue
            document.add_paragraph(line[2:].strip(), style="Heading 1")
            continue

        if line.startswith("## "):
            document.add_paragraph(line[3:].strip(), style="Heading 1")
            continue

        if line.startswith("### "):
            document.add_paragraph(line[4:].strip(), style="Heading 2")
            continue

        if line.startswith("#### "):
            document.add_paragraph(line[5:].strip(), style="Heading 3")
            continue

        image_match = re.match(r"^!\[(.+)\]\((.+)\)$", line.strip())
        if image_match:
            caption = image_match.group(1).strip()
            image_rel = image_match.group(2).strip()
            image_path = (markdown_path.parent / image_rel).resolve()
            add_image(document, image_path, caption)
            continue

        if re.match(r"^(Project|Published on|Version):", line):
            continue

        add_normal_paragraph(document, line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def main(argv: list[str]) -> int:
    if len(argv) != 3:
        print("Usage: generate_srs_doc.py <markdown_path> <output_docx_path>")
        return 1

    markdown_path = Path(argv[1])
    output_path = Path(argv[2])
    markdown_to_docx(markdown_path, output_path)
    print(f"Generated: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
