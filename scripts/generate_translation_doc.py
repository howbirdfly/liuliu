from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BODY_FONT = "宋体"
TITLE_FONT = "黑体"
EN_FONT = "Times New Roman"


def set_run_font(run, east_asia: str = BODY_FONT, ascii_font: str = EN_FONT, size: int = 12, bold: bool = False) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.bold = bold


def configure(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = EN_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(12)


def add_paragraph(document: Document, text: str, size: int = 12, bold: bool = False, center: bool = False, left_indent_pt: float = 0) -> None:
    p = document.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if left_indent_pt:
        p.paragraph_format.left_indent = Pt(left_indent_pt)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, east_asia=TITLE_FONT if bold and size >= 14 else BODY_FONT)


def add_list_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    r = p.add_run(text)
    set_run_font(r)


def render(md_path: Path, docx_path: Path) -> None:
    document = Document()
    configure(document)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    for line in lines:
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
            continue

        if stripped.startswith("# "):
            add_paragraph(document, stripped[2:].strip(), size=16, bold=True, center=True)
            continue
        if stripped.startswith("## "):
            add_paragraph(document, stripped[3:].strip(), size=14, bold=True)
            continue
        if stripped.startswith("### "):
            add_paragraph(document, stripped[4:].strip(), size=12, bold=True)
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            add_list_paragraph(document, f"{numbered.group(1)}. {numbered.group(2)}")
            continue

        dashed = re.match(r"^[-•]\s+(.*)$", stripped)
        if dashed:
            add_list_paragraph(document, f"• {dashed.group(1)}")
            continue

        add_paragraph(document, stripped)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(docx_path))


def main(argv: list[str]) -> int:
    if len(argv) != 3:
        print("Usage: generate_translation_doc.py <input.md> <output.docx>")
        return 1
    render(Path(argv[1]), Path(argv[2]))
    print(argv[2])
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
