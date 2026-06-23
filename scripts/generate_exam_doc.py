from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


BODY_FONT = "宋体"
HEADING_FONT = "黑体"
EN_FONT = "Times New Roman"


def set_run_font(run, east_asia: str = BODY_FONT, ascii_font: str = EN_FONT, size: int = 12, bold: bool = False) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.bold = bold


def configure_doc(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(90)
    section.right_margin = Pt(90)

    normal = document.styles["Normal"]
    normal.font.name = EN_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(12)


def add_title(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, east_asia=HEADING_FONT, size=16, bold=True)


def add_center_line(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=12)


def add_heading(document: Document, text: str, level: int) -> None:
    p = document.add_paragraph()
    r = p.add_run(text)
    size = 14 if level == 2 else 12
    set_run_font(r, east_asia=HEADING_FONT, size=size, bold=True)


def add_body(document: Document, text: str, indent: float = 0.0) -> None:
    p = document.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Pt(indent)
    r = p.add_run(text)
    set_run_font(r, size=12)


def render_markdown(markdown_path: Path, output_docx: Path) -> None:
    document = Document()
    configure_doc(document)

    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    first_heading = True

    for line in lines:
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
            continue

        if stripped.startswith("# "):
            if first_heading:
                add_title(document, stripped[2:].strip())
                first_heading = False
            else:
                add_heading(document, stripped[2:].strip(), 1)
            continue

        if stripped.startswith("## "):
            add_heading(document, stripped[3:].strip(), 2)
            continue

        if re.match(r"^(适用范围|考试形式|考试时间|试卷总分)：", stripped):
            add_center_line(document, stripped)
            continue

        add_body(document, stripped)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_docx))


def main(argv: list[str]) -> int:
    if len(argv) != 3:
        print("Usage: generate_exam_doc.py <markdown> <docx>")
        return 1
    render_markdown(Path(argv[1]), Path(argv[2]))
    print(argv[2])
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
